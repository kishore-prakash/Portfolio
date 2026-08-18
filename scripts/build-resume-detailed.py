#!/usr/bin/env python3
"""Build the detailed, human-facing resume from src/data/resume.json.

This is the counterpart to build-resume.py. That one is the two-page ATS
resume: single text flow, no tables, no colour, nothing a parser can trip on.
This one is read by people — a hiring manager, an interview panel, an internal
profile — so it trades parser-safety for design and completeness. It uses
tables, shading and colour, which means it is deliberately NOT ATS-safe. Send
the default build to job portals; send this one to humans.

Design language mirrors the website: Zinc ink on white, one blue accent
(#0055D4), Georgia for display type against Calibri for body copy, wide
letter-spaced section labels, and hairline rules instead of boxes.
"""

import argparse
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

sys.path.insert(0, str(__import__("pathlib").Path(__file__).resolve().parent))
from resume_common import ROOT, OUT_DIR, clean, group_by, has_placeholder, load  # noqa: E402

OUT = OUT_DIR / "Kishore-Prakash-Resume-Detailed.docx"

# Same palette as the site's light theme, so the document and kishoreprakash.in
# read as one thing.
INK = RGBColor(0x18, 0x18, 0x1B)
MUTED = RGBColor(0x52, 0x52, 0x5B)
FAINT = RGBColor(0x71, 0x71, 0x7A)
ACCENT = RGBColor(0x00, 0x55, 0xD4)
ACCENT_HEX = "0055D4"
SOFT_HEX = "EEF4FF"
SUBTLE_HEX = "F7F7F8"
LINE_HEX = "E4E4E7"

DISPLAY = "Georgia"
BODY = "Calibri"

CONTENT_WIDTH = Inches(7.3)
RIGHT_TAB = Inches(7.3)

CATEGORY_LABELS = {
    "enterprise": "Enterprise applications",
    "sdk": "SDKs and developer tooling",
    "nfc": "NFC and connectivity apps",
    "personal": "Independent products",
}


# --------------------------------------------------------------------------
# low-level docx helpers
# --------------------------------------------------------------------------

def style_run(run, *, font=BODY, size=10.5, bold=False, italic=False,
              color=INK, spacing=None, caps=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rPr.rFonts.set(qn("w:ascii"), font)
    rPr.rFonts.set(qn("w:hAnsi"), font)
    rPr.rFonts.set(qn("w:cs"), font)
    if spacing is not None:
        el = OxmlElement("w:spacing")          # letter-spacing, in twentieths of a point
        el.set(qn("w:val"), str(int(spacing * 20)))
        rPr.append(el)
    if caps:
        rPr.append(OxmlElement("w:caps"))
    return run


def text(container, content="", *, font=BODY, size=10.5, bold=False, italic=False,
         color=INK, spacing=None, caps=False, before=0, after=0, line=1.15,
         align=None, keep_with_next=False):
    p = container.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align
    if keep_with_next:
        pf.keep_with_next = True
    if content:
        style_run(p.add_run(content), font=font, size=size, bold=bold, italic=italic,
                  color=color, spacing=spacing, caps=caps)
    return p


def shade(element, hex_fill: str) -> None:
    """Background fill for a paragraph or a table cell."""
    pr = element.get_or_add_pPr() if element.tag.endswith("}p") else element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    pr.append(shd)


def rule(p, *, color=LINE_HEX, size=6, position="bottom", space=4) -> None:
    """A hairline instead of a box — the whole document is built from these."""
    pPr = p._p.get_or_add_pPr()
    borders = pPr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        pPr.append(borders)
    edge = OxmlElement(f"w:{position}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)
    borders.append(edge)


def cell_borders(cell, **edges) -> None:
    """edges: left=("0055D4", 18) etc. Anything unnamed is removed."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        if side in edges:
            color, size = edges[side]
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(size))
            el.set(qn("w:color"), color)
        else:
            el.set(qn("w:val"), "nil")
        el.set(qn("w:space"), "0")
        borders.append(el)
    tcPr.append(borders)


def cell_padding(table, *, top=60, left=110, bottom=60, right=110) -> None:
    """Table cell margins, in twentieths of a point — this is the card padding."""
    tblPr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)


def no_split(row) -> None:
    """Stop Word breaking a row — and so a card — across a page boundary."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def bare_table(doc, rows, cols, widths=None):
    t = doc.add_table(rows=rows, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    t.style = "Table Grid"
    for row in t.rows:
        for cell in row.cells:
            cell_borders(cell)
    if widths:
        for row in t.rows:
            for cell, w in zip(row.cells, widths):
                cell.width = w
    return t


def spacer(doc, points=6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    for run in p.runs:
        run.font.size = Pt(points)
    p.add_run().font.size = Pt(points)


def field(p, instruction: str, *, size=8.5, color=FAINT):
    """A Word field (PAGE, NUMPAGES) — recalculated on open, unlike literal text."""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f" {instruction} ")
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), BODY)
    fonts.set(qn("w:hAnsi"), BODY)
    rPr.append(fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    col = OxmlElement("w:color")
    col.set(qn("w:val"), f"{color}")
    rPr.append(col)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = "1"
    run.append(t)
    fld.append(run)
    p._p.append(fld)


# --------------------------------------------------------------------------
# document furniture
# --------------------------------------------------------------------------

def setup(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.15

    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.footer_distance = Inches(0.3)

    bullet = doc.styles["List Bullet"]
    bullet.font.name = BODY
    bullet.font.size = Pt(10.5)
    bullet.font.color.rgb = MUTED
    bpf = bullet.paragraph_format
    bpf.left_indent = Inches(0.24)
    bpf.first_line_indent = Inches(-0.15)
    bpf.space_after = Pt(3)
    bpf.line_spacing = 1.15


def footer(doc: Document, name: str) -> None:
    doc.styles["Footer"].paragraph_format.tab_stops.clear_all()
    p = doc.sections[0].footer.paragraphs[0]
    p.paragraph_format.tab_stops.clear_all()
    p.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
    style_run(p.add_run(f"{name}  ·  Detailed resume"), size=8.5, color=FAINT,
              spacing=0.3, caps=True)
    style_run(p.add_run("\t"), size=8.5, color=FAINT)
    field(p, "PAGE")
    style_run(p.add_run(" / "), size=8.5, color=FAINT)
    field(p, "NUMPAGES")


def section_label(doc, label: str) -> None:
    p = text(doc, label, size=9, bold=True, color=ACCENT, spacing=1.6, caps=True,
             before=17, after=5, keep_with_next=True)
    rule(p, color=ACCENT_HEX, size=8, space=3)


def lede(doc, content: str) -> None:
    text(doc, content, size=10, color=FAINT, italic=True, after=6, keep_with_next=True)


def entry_line(doc, left: str, right: str, *, font=DISPLAY, size=11.5, bold=True,
               color=INK, before=9, after=0, right_color=FAINT, right_size=9.5):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.keep_with_next = True
    pf.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
    style_run(p.add_run(left), font=font, size=size, bold=bold, color=color)
    if right:
        style_run(p.add_run("\t" + right), size=right_size, color=right_color, spacing=0.2)
    return p


def bullets(doc, items, *, color=MUTED):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        style_run(p.add_run(clean(item)), size=10.5, color=color)


def chips(doc, items, cols=3, gutter=0.1):
    """Soft-filled cells standing in for the website's chip row.

    Word tables have no gap between cells, so the gutters are unshaded spacer
    columns and each visual row is its own one-row table with a spacer after
    it. Without that, a chip grid reads as one solid block.
    """
    width = Inches((7.3 - gutter * (cols - 1)) / cols)
    widths = []
    for i in range(cols):
        if i:
            widths.append(Inches(gutter))
        widths.append(width)
    for start in range(0, len(items), cols):
        row = items[start:start + cols]
        t = bare_table(doc, 1, len(widths), widths=widths)
        cell_padding(t, top=50, bottom=50, left=100, right=100)
        no_split(t.rows[0])
        for i, w in enumerate(widths):
            cell = t.cell(0, i)
            cell.width = w
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1
            slot = i // 2
            if i % 2 == 0 and slot < len(row):
                shade(cell._tc, SOFT_HEX)
                style_run(p.add_run(row[slot]), size=9.5, color=INK)
        spacer(doc, 3)


def card(doc, *, accent=ACCENT_HEX, fill=SUBTLE_HEX):
    """One-cell table used as a card: soft fill, a single accent edge, no box."""
    t = bare_table(doc, 1, 1, widths=[CONTENT_WIDTH])
    cell_padding(t, top=90, bottom=90, left=140, right=140)
    no_split(t.rows[0])
    cell = t.cell(0, 0)
    cell.width = CONTENT_WIDTH
    shade(cell._tc, fill)
    cell_borders(cell, left=(accent, 18))
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    return cell


# --------------------------------------------------------------------------
# sections
# --------------------------------------------------------------------------

def masthead(doc, b: dict) -> None:
    text(doc, b["name"], font=DISPLAY, size=27, bold=True, after=1, line=1)
    text(doc, b["title"], size=10.5, bold=True, color=ACCENT, spacing=1.8, caps=True,
         after=7, line=1)
    p = text(doc, size=9.5, after=0, line=1.2)
    bits = [b["location"], b["phone"], b["email"]]
    style_run(p.add_run("   ·   ".join(bits)), size=9.5, color=MUTED)
    p2 = text(doc, size=9.5, after=0, line=1.2)
    links = [b["linkedinLabel"], b["websiteLabel"], b["githubLabel"]]
    style_run(p2.add_run("   ·   ".join(links)), size=9.5, color=MUTED)
    rule(p2, color=ACCENT_HEX, size=14, space=8)


def profile(doc, data: dict) -> None:
    section_label(doc, "Profile")
    text(doc, data["basics"]["summary"], size=11, color=MUTED, line=1.25, after=8)
    chips(doc, data["competencies"], cols=3)


def skills(doc, data: dict) -> None:
    section_label(doc, "Technical skills")
    for group in data["skills"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(1.65))
        style_run(p.add_run(group["group"]), size=9.5, bold=True, color=ACCENT,
                  caps=True, spacing=0.4)
        style_run(p.add_run("\t" + ", ".join(group["items"])), size=10.5, color=MUTED)


def experience(doc, data: dict, dropped: list[str]) -> None:
    section_label(doc, "Experience")
    for company in data["experience"]:
        entry_line(doc, company["company"], company["location"], before=11)
        if company.get("context"):
            text(doc, company["context"], size=9.5, italic=True, color=FAINT,
                 after=2, keep_with_next=True)
        for role in company["roles"]:
            entry_line(doc, role["title"], f"{role['start']} – {role['end']}",
                       font=BODY, size=10.5, color=ACCENT, before=5, after=3,
                       right_color=MUTED)
            for raw in role["bullets"]:
                if has_placeholder(raw):
                    dropped.append(f"{company['company']} / {role['title']}: {raw}")
            bullets(doc, role["bullets"])


def projects(doc, data: dict) -> None:
    section_label(doc, "Projects")
    lede(doc, f"{len(data['projects'])} shipped products and libraries, "
              f"grouped by type.")
    for category, label in CATEGORY_LABELS.items():
        items = [p for p in data["projects"] if p.get("category") == category]
        if not items:
            continue
        text(doc, label, font=DISPLAY, size=11, bold=True, color=INK,
             before=12, after=4, keep_with_next=True)
        for proj in items:
            cell = card(doc)
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(6.85), WD_TAB_ALIGNMENT.RIGHT)
            style_run(p.add_run(proj["name"]), font=DISPLAY, size=11, bold=True)
            style_run(p.add_run("\t" + (proj.get("years") or "")), size=9, color=FAINT)

            meta = [proj.get("org"), f"team of {proj['teamSize']}"]
            mp = cell.add_paragraph()
            mp.paragraph_format.space_after = Pt(4)
            style_run(mp.add_run("  ·  ".join(x for x in meta if x)),
                      size=9, color=FAINT, caps=True, spacing=0.3)

            bp = cell.add_paragraph()
            bp.paragraph_format.space_after = Pt(3)
            style_run(bp.add_run(proj["blurb"]), size=10, color=MUTED)

            for h in proj.get("highlights", []):
                hp = cell.add_paragraph()
                hp.paragraph_format.left_indent = Inches(0.16)
                hp.paragraph_format.space_after = Pt(1)
                style_run(hp.add_run("— "), size=10, color=ACCENT)
                style_run(hp.add_run(clean(h)), size=10, color=MUTED)

            tp = cell.add_paragraph()
            tp.paragraph_format.space_before = Pt(4)
            tp.paragraph_format.space_after = Pt(0)
            style_run(tp.add_run(", ".join(proj["tech"])), size=9, color=INK)
            if proj.get("role"):
                rp = cell.add_paragraph()
                rp.paragraph_format.space_after = Pt(0)
                style_run(rp.add_run("Role: "), size=9, bold=True, color=FAINT)
                style_run(rp.add_run("; ".join(proj["role"])), size=9, color=FAINT)
            lp = cell.add_paragraph()
            lp.paragraph_format.space_after = Pt(0)
            if proj.get("link"):
                style_run(lp.add_run(proj["link"]), size=9, color=ACCENT)
            elif proj.get("delisted"):
                style_run(lp.add_run("No longer on the App Store"), size=9,
                          italic=True, color=FAINT)
            spacer(doc, 5)


def awards(doc, data: dict) -> None:
    section_label(doc, "Awards and recognition")
    lede(doc, f"{len(data['awards'])} recognitions, newest first.")
    dated = [a for a in data["awards"] if a.get("date")]
    undated = [a for a in data["awards"] if not a.get("date")]
    for org, items in group_by(dated, lambda a: a["org"]):
        p = text(doc, org, font=DISPLAY, size=11, bold=True, before=10, after=3,
                 keep_with_next=True)
        p.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
        style_run(p.add_run("\t" + f"{len(items)} award{'s' if len(items) > 1 else ''}"),
                  size=9, color=FAINT, caps=True, spacing=0.3)
        rule(p, color=LINE_HEX, size=6, space=2)
        for a in items:
            entry_line(doc, a["name"], a.get("date", ""), font=BODY, size=10.5,
                       color=INK, before=6, after=1, right_color=FAINT)
            if a.get("detail"):
                text(doc, a["detail"], size=9.5, color=MUTED, after=1, line=1.2)

    if undated:
        p = text(doc, "Earlier recognition", font=DISPLAY, size=11, bold=True,
                 before=10, after=3, keep_with_next=True)
        p.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
        style_run(p.add_run("\t" + f"{len(undated)} awards"), size=9, color=FAINT,
                  caps=True, spacing=0.3)
        rule(p, color=LINE_HEX, size=6, space=2)
        t = bare_table(doc, -(-len(undated) // 2), 2,
                       widths=[Inches(3.65), Inches(3.65)])
        cell_padding(t, top=40, bottom=40, left=0, right=120)
        for i, a in enumerate(undated):
            cell = t.cell(i // 2, i % 2)
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_after = Pt(0)
            style_run(cp.add_run(a["name"]), size=10, color=INK)
            style_run(cp.add_run(f"  ·  {a['org']}"), size=9.5, color=FAINT)
        spacer(doc, 4)


def background(doc, data: dict) -> None:
    section_label(doc, "Education, certifications and languages")
    for e in data["education"]:
        entry_line(doc, e["degree"], f"{e['start']} – {e['end']}", font=BODY,
                   size=10.5, before=6, after=1)
        text(doc, f"{e['school']}  ·  {e['affiliation']}", size=9.5, color=MUTED, after=2)

    t = bare_table(doc, 1, 2, widths=[Inches(3.65), Inches(3.65)])
    cell_padding(t, top=60, bottom=60, left=0, right=140)
    for cell, label, items in (
        (t.cell(0, 0), "Certifications", [c["name"] for c in data["certifications"]]),
        (t.cell(0, 1), "Languages",
         [f"{l['name']} — {l['level'].lower()}" for l in data["languages"]]),
    ):
        cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
        lp = cell.add_paragraph()
        lp.paragraph_format.space_before = Pt(8)
        lp.paragraph_format.space_after = Pt(3)
        style_run(lp.add_run(label), size=9, bold=True, color=ACCENT, caps=True,
                  spacing=1.2)
        for item in items:
            ip = cell.add_paragraph()
            ip.paragraph_format.space_after = Pt(2)
            style_run(ip.add_run("· "), size=10, color=ACCENT)
            style_run(ip.add_run(item), size=10, color=MUTED)


def build(data: dict) -> tuple[Document, list[str]]:
    dropped: list[str] = []
    doc = Document()
    setup(doc)
    footer(doc, data["basics"]["name"])

    masthead(doc, data["basics"])
    profile(doc, data)
    skills(doc, data)
    experience(doc, data, dropped)
    projects(doc, data)
    awards(doc, data)
    background(doc, data)
    return doc, dropped


def main() -> int:
    argparse.ArgumentParser(description=__doc__).parse_args()
    data = load()
    doc, dropped = build(data)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT.relative_to(ROOT)}")
    if dropped:
        print(f"\n{len(dropped)} bullet(s) had unconfirmed metric placeholders. "
              f"The placeholder clause was removed and the bullet shipped "
              f"qualitatively — no number was invented:\n")
        for d in dropped:
            print(f"  - {d}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
