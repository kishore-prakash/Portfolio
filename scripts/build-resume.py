#!/usr/bin/env python3
"""Generate an ATS-friendly .docx resume from src/data/resume.json.

ATS parsers choke on tables, text boxes, columns, headers/footers and images,
so this builds a single linear text flow with standard section headings and a
real Word list style for bullets. Dates are placed with a right-aligned tab
stop rather than a table cell, which keeps them on the same visual line while
still parsing as one paragraph.
"""

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

ROOT = Path(__file__).resolve().parent.parent
DATA = ROOT / "src" / "data" / "resume.json"
OUT = ROOT / "public" / "Kishore-Prakash-Resume.docx"

BODY_FONT = "Calibri"
BODY_SIZE = Pt(10)
NAME_SIZE = Pt(20)
HEADING_SIZE = Pt(12)
INK = RGBColor(0x1A, 0x1A, 0x1A)
RULE = "1F4E79"

# Metrics still awaiting confirmation are written as [[...]] in resume.json.
# Rather than shipping placeholder text or inventing a number, strip the
# placeholder clause and keep the sentence qualitative.
# A placeholder sitting between two commas collapses to one comma, so the
# sentence still reads as a list rather than losing its punctuation.
PLACEHOLDER_MID = re.compile(r"\s*,\s*\[\[[^\]]*\]\]\s*,\s*")
PLACEHOLDER = re.compile(r"\s*,?\s*\[\[[^\]]*\]\]")


def clean(text: str) -> str:
    text = PLACEHOLDER_MID.sub(", ", text)
    text = PLACEHOLDER.sub("", text)
    text = re.sub(r"\s+([.,;])", r"\1", text)
    text = re.sub(r"\s{2,}", " ", text).strip()
    if text and text[-1] not in ".!?":
        text += "."
    return text


def has_placeholder(text: str) -> bool:
    return "[[" in text


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.font.color.rgb = INK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    bullet = doc.styles["List Bullet"]
    bullet.font.name = BODY_FONT
    bullet.font.size = BODY_SIZE
    bullet.font.color.rgb = INK
    bpf = bullet.paragraph_format
    bpf.left_indent = Inches(0.22)
    bpf.first_line_indent = Inches(-0.15)
    bpf.space_after = Pt(1.5)
    bpf.line_spacing = 1.0


def para(doc, text="", *, bold=False, italic=False, size=None, space_before=0,
         space_after=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = size
    return p


def bottom_rule(p) -> None:
    """A single bottom border under a heading. Borders parse fine; tables don't."""
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), RULE)
    borders.append(bottom)
    pPr.append(borders)


def heading(doc, text: str) -> None:
    p = para(doc, text.upper(), bold=True, size=HEADING_SIZE,
             space_before=9, space_after=3)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    bottom_rule(p)


def role_line(doc, left: str, right: str, *, bold_left=True, italic_left=False):
    """Title on the left, dates flush right, via a right tab stop."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    tab_pos = Inches(7.2)
    p.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT)
    run = p.add_run(left)
    run.bold = bold_left
    run.italic = italic_left
    p.add_run("\t" + right)
    return p


def bullet(doc, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(1.5)


def build(data: dict) -> tuple[Document, list[str], list[str]]:
    dropped: list[str] = []
    trimmed: list[str] = []
    doc = Document()
    setup_styles(doc)
    b = data["basics"]

    # --- Contact block: plain text lines, never a header or a table. ---
    name = para(doc, b["name"], bold=True, size=NAME_SIZE, align=WD_ALIGN_PARAGRAPH.CENTER)
    name.paragraph_format.space_after = Pt(1)
    para(doc, b["title"], size=Pt(11.5), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    para(doc, f"{b['location']}  |  {b['phone']}  |  {b['email']}",
         size=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    para(doc, f"{b['linkedinLabel']}  |  {b['websiteLabel']}  |  {b['githubLabel']}",
         size=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER)

    heading(doc, "Professional Summary")
    para(doc, b["summary"], space_after=2)

    heading(doc, "Core Competencies")
    comps = data["competencies"]
    half = (len(comps) + 1) // 2
    for row in (comps[:half], comps[half:]):
        para(doc, "  •  ".join(row), space_after=1)

    heading(doc, "Technical Skills")
    # Merged into six lines so the skills block stays scannable and the
    # resume stays inside two pages.
    by_group = {g["group"]: g["items"] for g in data["skills"]}
    merged = [
        ("Languages", by_group["Languages"]),
        ("Platforms", by_group["Apple Platforms"] + by_group["Cross-Platform"]),
        ("Connectivity", by_group["Connectivity"]),
        ("Architecture & APIs", by_group["Architecture & APIs"]),
        ("CI/CD & Testing", by_group["CI/CD"] + by_group["Testing"]),
        ("Data & Tooling", by_group["Data"] + by_group["Documentation"] + by_group["Tooling"]),
    ]
    for label, items in merged:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.add_run(f"{label}: ").bold = True
        p.add_run(", ".join(items))

    heading(doc, "Professional Experience")
    # Recency-weighted: the current and recent roles carry the detail, older
    # ones are trimmed to their headline achievements. Standard practice, and
    # it is what keeps this inside two pages.
    BULLET_CAP = [5, 4, 3, 3, 3, 2, 2]
    role_index = 0
    for company in data["experience"]:
        p = para(doc, space_before=6, space_after=0)
        run = p.add_run(f"{company['company']}")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(f" — {company['location']}")
        for role in company["roles"]:
            role_line(doc, role["title"], f"{role['start']} – {role['end']}", italic_left=True)
            cap = BULLET_CAP[min(role_index, len(BULLET_CAP) - 1)]
            for raw in role["bullets"][:cap]:
                if has_placeholder(raw):
                    dropped.append(f"{company['company']} / {role['title']}: {raw}")
                bullet(doc, clean(raw))
            if len(role["bullets"]) > cap:
                trimmed.append(
                    f"{company['company']} / {role['title']}: "
                    f"{len(role['bullets']) - cap} bullet(s) held back for length"
                )
            role_index += 1

    heading(doc, "Selected Projects")
    featured = [p for p in data["projects"] if p.get("resumeFeature")]
    rest = [p for p in data["projects"] if not p.get("resumeFeature")]

    for proj in featured:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        p.add_run(proj["name"]).bold = True
        meta = [proj.get("org"), proj.get("years"), f"team of {proj['teamSize']}"]
        p.add_run(" — " + " | ".join(x for x in meta if x)).italic = True

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        first_sentence = proj["blurb"].split(". ")[0].rstrip(".") + ". "
        p.add_run(first_sentence)
        p.add_run(", ".join(proj["tech"]) + ".").italic = True

    if rest:
        # Everything else as one compact run-on paragraph rather than a bullet
        # each: keeps the breadth visible without spending half a page on it.
        grouped: dict[str, list[str]] = {}
        for proj in rest:
            grouped.setdefault(proj.get("org") or "Other", []).append(proj["name"])
        chunks = [
            f"{', '.join(names)} ({org})" if org != "Other" else ", ".join(names)
            for org, names in grouped.items()
        ]
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(1)
        p.add_run("Also shipped: ").bold = True
        p.add_run("; ".join(chunks) + ".")

    heading(doc, "Education")
    for e in data["education"]:
        role_line(doc, e["degree"], f"{e['start']} – {e['end']}")
        para(doc, f"{e['school']} | {e['affiliation']}", size=Pt(10))

    heading(doc, "Certifications, Awards & Languages")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.add_run("Certifications: ").bold = True
    p.add_run("; ".join(c["name"] for c in data["certifications"]) + ".")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.add_run("Awards: ").bold = True
    p.add_run("; ".join(f"{a['name']} ({a['org']})" for a in data["awards"]) + ".")

    p = doc.add_paragraph()
    p.add_run("Languages: ").bold = True
    p.add_run("; ".join(f"{l['name']} — {l['level'].lower()}" for l in data["languages"]) + ".")

    return doc, dropped, trimmed


def main() -> int:
    data = json.loads(DATA.read_text())
    doc, dropped, trimmed = build(data)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT.relative_to(ROOT)}")

    if dropped:
        print(f"\n{len(dropped)} bullet(s) had unconfirmed metric placeholders. "
              f"The placeholder clause was removed and the bullet shipped "
              f"qualitatively — no number was invented:\n")
        for d in dropped:
            print(f"  - {d}")
    if trimmed:
        print(f"\nTrimmed for length (still in resume.json and on the website):\n")
        for t in trimmed:
            print(f"  - {t}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
